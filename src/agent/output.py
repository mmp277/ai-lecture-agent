from typing import List, Tuple, Dict
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import subprocess
import shutil


def write_summary_docx(output_path: str, file_summaries: List[Tuple[str, List[str]]]) -> None:
	doc = Document()
	doc.add_heading("Lecture Summaries", level=1)
	for filename, summaries in file_summaries:
		doc.add_heading(filename, level=2)
		for s in summaries:
			p = doc.add_paragraph(s)
			p_format = p.paragraph_format
			p_format.space_after = Pt(6)
		doc.add_page_break()
	doc.save(output_path)


def write_flashcards_docx(output_path: str, flashcards_by_file: List[Tuple[str, List[Tuple[str, str]]]]) -> None:
	doc = Document()
	doc.add_heading("Flashcards", level=1)
	for filename, cards in flashcards_by_file:
		doc.add_heading(filename, level=2)
		for idx, (q, a) in enumerate(cards, start=1):
			q_p = doc.add_paragraph()
			run_q = q_p.add_run(f"Q{idx}. {q}")
			run_q.bold = True
			a_p = doc.add_paragraph(f"A{idx}. {a}")
			a_p_format = a_p.paragraph_format
			a_p_format.space_after = Pt(6)
		doc.add_page_break()
	doc.save(output_path)


def write_formula_sheet_docx(output_path: str, formulas_by_file: List[Tuple[str, List[Tuple[str, Dict[str, str]]]]]) -> None:
	doc = Document()
	doc.add_heading("Formula Sheet", level=1)
	for filename, items in formulas_by_file:
		doc.add_heading(filename, level=2)
		for eq, defs in items:
			p_eq = doc.add_paragraph()
			r = p_eq.add_run(eq)
			r.italic = True
			p_eq.alignment = WD_ALIGN_PARAGRAPH.CENTER

			if defs:
				doc.add_paragraph("Variables:")
				for sym, desc in defs.items():
					doc.add_paragraph(f"{sym}: {desc}", style=None)
			doc.add_paragraph("")
		doc.add_page_break()
	doc.save(output_path)


# ---- LaTeX writer for high-quality formula rendering ----
LATEX_PREAMBLE = r"""\documentclass[11pt]{article}
\usepackage[margin=1in]{geometry}
\usepackage{amsmath, amssymb}
\usepackage{array}
\usepackage{booktabs}
\usepackage{hyperref}
\usepackage{longtable}
\title{Formula Sheet}
\date{}
\begin{document}
\maketitle
"""

LATEX_END = "\\end{document}\n"


def _latex_escape(text: str) -> str:
	return (
		text.replace("\\", r"\textbackslash{}")
		.replace("_", r"\_")
		.replace("%", r"\%")
		.replace("#", r"\#")
		.replace("&", r"\&")
		.replace("{", r"\{")
		.replace("}", r"\}")
	)


def write_formula_sheet_latex(output_tex_path: str, formulas_by_file: List[Tuple[str, List[Tuple[str, Dict[str, str]]]]], compile_pdf: bool = True) -> None:
	os.makedirs(os.path.dirname(output_tex_path) or ".", exist_ok=True)
	lines: List[str] = [LATEX_PREAMBLE]
	for filename, items in formulas_by_file:
		lines.append(f"\\section*{{{_latex_escape(filename)}}}\n")
		for eq, defs in items:
			# eq is already formatted with $ ... $ or \[ ... \]
			if eq.strip().startswith("$"):
				tex_eq = eq.strip().strip("$").strip()
				lines.append("\\[ " + tex_eq + " \\]\n")
			elif eq.strip().startswith("\\["):
				lines.append(eq + "\n")
			else:
				lines.append("\\[ " + _latex_escape(eq) + " \\]\n")

			if defs:
				lines.append("\\noindent\\textbf{Variables}\\:\\")
				lines.append("\\begin{longtable}{@{}p{0.18\\textwidth}p{0.75\\textwidth}@{}}\n\\toprule\n\\textbf{Symbol} & \\textbf{Meaning} \\\\ \\midrule\n")
				for sym, desc in defs.items():
					lines.append(f"{_latex_escape(sym)} & {_latex_escape(desc)} \\\\ \n")
				lines.append("\\bottomrule\n\\end{longtable}\n")
			lines.append("\n")
	lines.append(LATEX_END)

	with open(output_tex_path, "w", encoding="utf-8") as f:
		f.write("".join(lines))

	if compile_pdf and shutil.which("pdflatex"):
		cwd = os.path.dirname(output_tex_path) or "."
		tex_name = os.path.basename(output_tex_path)
		try:
			# Run pdflatex twice for stable refs
			subprocess.run(["pdflatex", "-interaction=nonstopmode", tex_name], cwd=cwd, check=True, stdout=subprocess.PIPE, stderr=subprocess.STDOUT)
			subprocess.run(["pdflatex", "-interaction=nonstopmode", tex_name], cwd=cwd, check=True, stdout=subprocess.PIPE, stderr=subprocess.STDOUT)
		except Exception:
			# If compilation fails, leave the .tex for manual compile
			pass
